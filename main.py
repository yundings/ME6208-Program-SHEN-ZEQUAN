#!/usr/bin/env python3
"""
AI Blog Generator with UI (Cover Image Generated Based on Title) - Bilingual Comments
=======================================================================================

本程序根据输入的英文博客主题生成 3 套不同的博客结果。每套结果包括：
  1. 博客标题（使用 GPT-3.5-turbo 生成，并带有结果集编号确保独特）
  2. 博客正文（包含引言、正文和结尾，使用 GPT-3.5-turbo 生成，并带有编号）
  3. 封面图片（使用 GPT-4o（通过 DALL·E-3 接口）生成），要求图片背景简单优雅，
     并在图片中央以大字显示生成的博客标题（title）——即按 title 生成图片
  4. 注释说明（使用 GPT-3.5-turbo 生成，描述该结果的独特创意特点）
所有结果整合保存为 Word 文档（默认文件名 "generated_blog.docx"）。
程序提供简单的 Tkinter UI 界面，用户无需命令行操作即可使用。

Dependencies:
  - openai
  - python-docx
  - requests
  - Pillow
  - Tkinter (通常内置于 Python)

请妥善保管你的 API 密钥。
"""

import os
import time
import threading
import openai
import requests
import tkinter as tk
from tkinter import scrolledtext, messagebox
from docx import Document
from docx.shared import Inches

# -------------------------------------------------------------------
# API Key 设置 / API Key Configuration
# -------------------------------------------------------------------
openai.api_key = ("Please input your API key here")


# 请确保上述 API Key 正确且安全

# -------------------------------------------------------------------
# UI 日志输出辅助函数 / Append Log to UI
# -------------------------------------------------------------------
def append_log(message: str):
    """
    将日志消息追加到 UI 的滚动文本框中。
    Append log messages to the UI's scrolled text widget.
    """
    global log_text
    if log_text:
        log_text.insert(tk.END, message)
        log_text.see(tk.END)


# -------------------------------------------------------------------
# 通用文本生成函数 / General Text Generation Function (GPT-3.5-turbo)
# -------------------------------------------------------------------
def generate_text(prompt: str, max_tokens: int) -> str:
    """
    使用 OpenAI ChatCompletion 接口生成文本，基于 GPT-3.5-turbo 模型。
    Generate text using OpenAI's ChatCompletion API with GPT-3.5-turbo.

    Args:
      prompt (str): 输入提示 / Input prompt.
      max_tokens (int): 最大 token 数限制 / Maximum tokens.

    Returns:
      result (str): 生成的文本 / Generated text.
    """
    payload = {
        "model": "gpt-3.5-turbo",
        "messages": [
            {"role": "system", "content": "You are a helpful text generation assistant."},
            {"role": "user", "content": prompt}
        ],
        "max_tokens": max_tokens,
        "temperature": 0.7
    }
    while True:
        try:
            response = openai.ChatCompletion.create(**payload)
            result = response["choices"][0]["message"]["content"].strip()
            return result
        except openai.error.RateLimitError:
            append_log("Text generation rate limit reached. Waiting 10 seconds...\n")
            time.sleep(10)
        except Exception as e:
            append_log(f"Error generating text: {e}. Retrying in 10 seconds...\n")
            time.sleep(10)


# -------------------------------------------------------------------
# 生成博客标题 / Generate Blog Title
# -------------------------------------------------------------------
def generate_blog_title(topic: str, set_no: int) -> str:
    """
    根据博客主题生成一个独特的英文博客标题，并加入结果集编号保证独特性。
    Generate a unique blog title based on the given topic.
    Include the set number to ensure each title is unique.

    Args:
      topic (str): 博客主题 / Blog topic.
      set_no (int): 结果集编号 / Set number.

    Returns:
      title (str): 生成的博客标题 / Generated blog title.
    """
    prompt = (
        f"Please generate an attractive blog title for the following topic: '{topic}'. "
        f"This is for result set {set_no}, so ensure it is unique and creative. "
        "Requirements: short, creative, and captivating."
    )
    return generate_text(prompt, max_tokens=50)


# -------------------------------------------------------------------
# 生成博客正文 / Generate Blog Content
# -------------------------------------------------------------------
def generate_blog_content(topic: str, set_no: int) -> str:
    """
    根据博客主题生成一篇完整的英文博客文章（包含引言、正文和结尾），并加入结果集编号确保每组内容独特。
    Generate an engaging English blog post on the given topic including introduction, body, and conclusion.
    Include the set number for uniqueness.

    Args:
      topic (str): 博客主题 / Blog topic.
      set_no (int): 结果集编号 / Set number.

    Returns:
      content (str): 生成的博客正文 / Generated blog content.
    """
    prompt = (
        f"Write an engaging English blog post about '{topic}' for result set {set_no} that includes an introduction, "
        "a main body, and a conclusion. The post should be between 500 to 800 words, logically structured, "
        "and include personal opinions. Ensure it is uniquely different from other sets."
    )
    return generate_text(prompt, max_tokens=1500)


# -------------------------------------------------------------------
# 生成注释说明 / Generate Annotation
# -------------------------------------------------------------------
def generate_annotation(topic: str, title: str, set_no: int) -> str:
    """
    根据博客主题和生成的标题生成一段注释说明，描述该结果的独特创意特点。
    Generate a brief annotation describing the unique creative features of this blog result.

    Args:
      topic (str): 博客主题 / Blog topic.
      title (str): 生成的博客标题 / Generated blog title.
      set_no (int): 结果集编号 / Set number.

    Returns:
      annotation (str): 生成的注释说明 / Generated annotation.
    """
    prompt = (
        f"Given the blog title '{title}' for the topic '{topic}', please provide a one-sentence annotation "
        f"describing the unique creative features of this blog result (Set {set_no})."
    )
    return generate_text(prompt, max_tokens=50)


# -------------------------------------------------------------------
# 使用 GPT-4o 生成封面图片（基于标题）/ Generate Cover Image using GPT-4o based on Title
# -------------------------------------------------------------------
def generate_cover_image_for_title(title: str, set_no: int) -> str:
    """
    使用 GPT-4o 图像生成接口生成一张 1024x1024 封面图片，要求图片背景简单优雅，
    并在中央以大字显示生成的博客标题（title）。在提示中加入结果集编号以确保背景设计各异，
    但显示的文字始终为生成的 title。

    Generate a 1024x1024 cover image using GPT-4o (via DALL·E-3) where the image features a simple,
    elegant background and displays the blog title (provided by GPT-3.5) in large, bold letters in the center.
    Include the set number in the prompt to vary the background design, ensuring the displayed text is exactly the title.

    Args:
      title (str): 生成的博客标题 / The blog title to display.
      set_no (int): 结果集编号 / Set number.

    Returns:
      image_url (str): 生成图片的 URL / URL of the generated image.
    """
    prompt = (
        f"Generate a 1024x1024 cover image with a simple, elegant background pattern. "
        f"In the center, display the following text in large, bold letters: '{title}'. "
        f"This is for variant {set_no}; vary the background style (different colors or patterns) "
        "but ensure the displayed text is exactly the given title."
    )
    while True:
        try:
            response = openai.Image.create(
                model="dall-e-3",  # 使用 DALL·E-3 接口生成图片，GPT-4o 可通过此接口调用
                prompt=prompt,
                n=1,
                size="1024x1024",
                response_format="url"
            )
            image_url = response["data"][0]["url"]
            return image_url
        except openai.error.RateLimitError:
            append_log("Image generation rate limit reached. Waiting 10 seconds...\n")
            time.sleep(10)
        except Exception as e:
            append_log(f"Error generating image with GPT-4o: {e}. Retrying in 10 seconds...\n")
            time.sleep(10)


# -------------------------------------------------------------------
# 下载图片 / Download Image
# -------------------------------------------------------------------
def download_image(image_url: str) -> str:
    """
    根据图片 URL 下载图片，并保存为临时文件。
    Download the image from the given URL and save it as a temporary file.

    Args:
      image_url (str): 图片的 URL / Image URL.

    Returns:
      temp_filename (str): 临时图片文件的路径 / Local path of the downloaded image.
                           若下载失败则返回空字符串 / Returns empty string on failure.
    """
    try:
        resp = requests.get(image_url)
        if resp.status_code == 200:
            temp_filename = f"temp_cover_image_set_{int(time.time())}.jpg"
            with open(temp_filename, "wb") as f:
                f.write(resp.content)
            return temp_filename
        else:
            append_log(f"Failed to download image, status code: {resp.status_code}\n")
            return ""
    except Exception as e:
        append_log(f"Error downloading image: {e}\n")
        return ""


# -------------------------------------------------------------------
# 保存到 Word 文档 / Save Blog Results to Word Document
# -------------------------------------------------------------------
def save_blog_to_word_multiple(sets: list, filename: str = "generated_blog.docx"):
    """
    将多个博客结果集（包括标题、注释、正文、封面图片）整合保存到一个 Word 文档中。
    Consolidate multiple blog result sets into a single Word document.

    Args:
      sets (list): 每个元素为一个字典，包含 'title', 'annotation', 'content', 'cover_image_path'.
                   / Each element is a dict with keys: 'title', 'annotation', 'content', 'cover_image_path'.
      filename (str): 保存的 Word 文件名称 / The filename for the Word document.
    """
    doc = Document()
    doc.add_heading("Generated Blog Results", level=1)
    for i, result in enumerate(sets, start=1):
        full_title = f"Set {i}: {result['title']} - {result['annotation']}"
        doc.add_heading(full_title, level=2)
        doc.add_paragraph(result["content"])
        if result["cover_image_path"] and os.path.exists(result["cover_image_path"]):
            doc.add_heading("Cover Image:", level=3)
            doc.add_picture(result["cover_image_path"], width=Inches(6))
        doc.add_paragraph("")  # 添加空行 / Add an empty paragraph for separation.
    os.makedirs(os.path.dirname(filename) if os.path.dirname(filename) else '.', exist_ok=True)
    doc.save(filename)
    append_log(f"Blog results have been saved to {filename}\n")


# -------------------------------------------------------------------
# 主生成流程 / Main Generation Flow (Triggered by UI)
# -------------------------------------------------------------------
def start_generation():
    """
    从 UI 获取博客主题，并依次生成 3 套博客结果：
      1. 生成博客标题 (GPT-3.5-turbo, 加入编号确保独特)
      2. 生成对应博客正文 (GPT-3.5-turbo)
      3. 根据生成的标题生成封面图片 (使用 GPT-4o, 图片中央显示标题)
      4. 生成注释说明 (GPT-3.5-turbo)
    最后将所有结果整合保存至 Word 文档中。

    Retrieve the blog topic from the UI, then sequentially generate three blog result sets:
      1. Generate blog title.
      2. Generate blog content.
      3. Generate cover image based on the generated title (using GPT-4o).
      4. Generate annotation.
    Finally, consolidate all results into a Word document.
    """
    topic = topic_entry.get().strip()
    if not topic:
        messagebox.showerror("Input Error", "Blog topic cannot be empty!")
        return

    generate_button.config(state=tk.DISABLED)  # 禁用按钮 / Disable button
    append_log(f"Starting generation for topic: {topic}\n")

    def worker():
        results = []
        for i in range(1, 4):
            append_log(f"\n----- Generating Set {i} -----\n")

            # 生成标题 / Generate Title
            title = generate_blog_title(topic, i)
            append_log(f"[Set {i}] Title: {title}\n")

            # 生成正文 / Generate Content
            content = generate_blog_content(topic, i)
            append_log(f"[Set {i}] Content generated.\n")

            # 生成封面图片（使用生成的标题作为图片显示文字） / Generate Cover Image based on Title
            image_url = generate_cover_image_for_title(title, i)
            append_log(f"[Set {i}] Cover image URL: {image_url}\n")

            cover_image_path = download_image(image_url)
            if cover_image_path:
                append_log(f"[Set {i}] Cover image downloaded: {cover_image_path}\n")
            else:
                append_log(f"[Set {i}] Cover image download failed.\n")

            # 生成注释说明 / Generate Annotation
            annotation = generate_annotation(topic, title, i)
            append_log(f"[Set {i}] Annotation: {annotation}\n")

            results.append({
                "title": title,
                "content": content,
                "cover_image_path": cover_image_path,
                "annotation": annotation
            })
            time.sleep(2)  # 等待避免频繁调用 / Pause briefly

        save_blog_to_word_multiple(results)
        append_log("All blog results generated and saved.\n")
        messagebox.showinfo("Success", "Blog results generated and saved successfully!")
        generate_button.config(state=tk.NORMAL)

    threading.Thread(target=worker, daemon=True).start()


# -------------------------------------------------------------------
# Tkinter UI 构建 / Build Tkinter UI
# -------------------------------------------------------------------
from tkinter import scrolledtext

root = tk.Tk()
root.title("AI Blog Generator with UI - English Version")
root.geometry("700x550")

topic_label = tk.Label(root, text="Enter Blog Topic (e.g., The Impact of AI on Digital Marketing):")
topic_label.pack(pady=10)

topic_entry = tk.Entry(root, width=80)
topic_entry.pack(pady=5)

generate_button = tk.Button(root, text="Generate Blog", command=start_generation)
generate_button.pack(pady=10)

log_text = scrolledtext.ScrolledText(root, width=80, height=20)
log_text.pack(pady=10)

#start gui
root.mainloop()
